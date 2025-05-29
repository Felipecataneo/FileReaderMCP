#!/usr/bin/env python3
"""
MCP Server para leitura e extração de conteúdo de arquivos
Suporta múltiplos formatos: PDF, DOCX, TXT, CSV, JSON, XML, HTML, etc.
Versão com suporte a URLs e processamento otimizado
"""

import asyncio
import json
import logging
import mimetypes
import os
import sys
import traceback
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Union, Tuple
import tempfile
import base64
from urllib.parse import urlparse
import urllib.request
import ssl # Import ssl

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("file_reader_mcp")

# Bibliotecas para diferentes tipos de arquivo
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    PyPDF2 = None
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    pdfplumber = None
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    Document = None
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    pd = None
    HAS_PANDAS = False

try:
    from bs4 import BeautifulSoup
    HAS_BEAUTIFULSOUP = True
except ImportError:
    BeautifulSoup = None
    HAS_BEAUTIFULSOUP = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    openpyxl = None
    HAS_OPENPYXL = False

try:
    import xml.etree.ElementTree as ET
    HAS_XML = True
except ImportError:
    ET = None
    HAS_XML = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    requests = None
    HAS_REQUESTS = False

# MCP imports - Use try-except for robustness
try:
    from mcp.server.fastmcp import FastMCP
    from mcp.types import (
        Resource,
        Tool,
        TextContent,
        ImageContent,
        EmbeddedResource,
        LoggingLevel
    )
    HAS_MCP = True
except ImportError as e:
    FastMCP = None
    Resource = Tool = TextContent = ImageContent = EmbeddedResource = LoggingLevel = None
    HAS_MCP = False
    # Log error later after logger is configured

# --- Adição para suporte a .env e SSE ---
try:
    from dotenv import load_dotenv
    HAS_DOTENV = True
except ImportError:
    load_dotenv = None
    HAS_DOTENV = False

if HAS_DOTENV:
    # Carrega variáveis de ambiente do arquivo .env. 
    # Adapte o caminho se o seu .env não estiver na mesma pasta do script ou na pasta pai.
    load_dotenv() 
    logger.info("dotenv loaded.")
else:
    logger.warning("python-dotenv not found. Host and port might need to be hardcoded or set via other means.")
# --- Fim da adição ---



# Log library status after logger is defined
logger.info(f"Library status - PyPDF2: {HAS_PYPDF2}, pdfplumber: {HAS_PDFPLUMBER}, python-docx: {HAS_DOCX}, pandas: {HAS_PANDAS}, BeautifulSoup: {HAS_BEAUTIFULSOUP}, openpyxl: {HAS_OPENPYXL}, XML: {HAS_XML}, requests: {HAS_REQUESTS}, MCP: {HAS_MCP}, dotenv: {HAS_DOTENV}")
if not HAS_MCP:
     logger.error("MCP libraries not available. The server will not be able to register tools or run.")


# Instância do FastMCP Server - Only create if MCP is available
if HAS_MCP:
    # --- Modificação para SSE: Adicionar host e port ---
    # Use variáveis de ambiente, ou defaults se não definidas
    MCP_HOST = os.getenv("MCP_HOST", "0.0.0.0") # Define o host (0.0.0.0 para aceitar conexões de qualquer IP)
    MCP_PORT = int(os.getenv("PORT", 8050)) # Define a porta
    
    mcp = FastMCP(
        name="file-reader",
        version="1.1.0",
        host=MCP_HOST,
        port=MCP_PORT,
    )
    logger.info(f"FastMCP instance created with host={MCP_HOST}, port={MCP_PORT}.")
else:
    mcp = None
    logger.warning("FastMCP instance could not be created.")


class FileReader:
    """Classe principal para leitura de arquivos"""

    def __init__(self):
        self.supported_extensions = {
            '.txt': self._read_text,
            '.md': self._read_text,
            '.py': self._read_text,
            '.js': self._read_text,
            '.html': self._read_html,
            '.htm': self._read_html,
            '.xml': self._read_xml,
            '.json': self._read_json,
            '.csv': self._read_csv,
            '.xlsx': self._read_excel,
            '.xls': self._read_excel,
            '.pdf': self._read_pdf,
            '.docx': self._read_docx,
            # Note: .doc requires extra libs like python-mammoth or subprocess calls to soffice/abiword
            # We'll keep .doc mapped to _read_docx, but it might fail for older .doc files
            '.doc': self._read_docx,
        }

        # Remove support for extensions if required libraries are missing
        if not HAS_PYPDF2 and not HAS_PDFPLUMBER:
             if '.pdf' in self.supported_extensions: del self.supported_extensions['.pdf']
             logger.warning("PDF support disabled: Neither PyPDF2 nor pdfplumber found.")
        if not HAS_DOCX:
             if '.docx' in self.supported_extensions: del self.supported_extensions['.docx']
             if '.doc' in self.supported_extensions: del self.supported_extensions['.doc']
             logger.warning("DOCX support disabled: python-docx not found.")
        if not HAS_PANDAS:
             if '.csv' in self.supported_extensions: del self.supported_extensions['.csv']
             if '.xlsx' in self.supported_extensions: del self.supported_extensions['.xlsx']
             if '.xls' in self.supported_extensions: del self.supported_extensions['.xls']
             logger.warning("CSV/Excel support disabled: pandas not found.")
        # HTML/XML can still be read as raw text even without parsing libraries, so keep them
        if not HAS_BEAUTIFULSOUP:
             logger.warning("HTML parsing (BeautifulSoup) not available.")
        if not HAS_XML:
             logger.warning("XML parsing (ElementTree) not available.")

        # Configuração para downloads
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.timeout = 30  # 30 segundos

        logger.info(f"FileReader inicializado. Formatos suportados: {list(self.supported_extensions.keys())}")


    def is_url(self, path: str) -> bool:
        """Verifica se o caminho é uma URL"""
        try:
            result = urlparse(path)
            # Check if scheme is http or https (or similar) and netloc exists
            return result.scheme in ['http', 'https', 'ftp', 'ftps'] and bool(result.netloc)
        except ValueError:
            return False


    async def download_file(self, url: str) -> Optional[Tuple[str, Optional[str]]]: # <- Change return type hint
        """Baixa arquivo de URL para arquivo temporário e retorna o Content-Type"""
        temp_path = None
        content_type = None # <- Initialize content_type
        try:
            logger.info(f"Baixando arquivo de: {url}")
            start_time = time.time()

            # ... (restante do código de SSL context, filename, suffix) ...
             # Extrai nome do arquivo da URL
            parsed_url = urlparse(url)
            filename = os.path.basename(parsed_url.path) or "downloaded_file"
            # Guess suffix from filename or MIME type if available
            # Ensure mimetypes.guess_type result is handled for potential None[0]
            mime_type_tuple = mimetypes.guess_type(url)
            mime_type_guess = mime_type_tuple[0] if mime_type_tuple else None # Safely get mime type string
            guessed_suffix_from_mime = mimetypes.guess_extension(mime_type_guess) if mime_type_guess else None # Safely get extension
            suffix = Path(filename).suffix or guessed_suffix_from_mime or '.tmp'
            # Add .html/.htm suffix if mime_type_guess looks like html, as a hint for OS mimetypes later IF Content-Type is not used
            # But using Content-Type directly is better. Let's remove this complex suffix logic and rely on Content-Type.
            # Keep it simple: suffix from path or just .tmp
            suffix = Path(filename).suffix or '.tmp'


            # Create temporary file first to get the path
            temp_file_obj = tempfile.NamedTemporaryFile(delete=False, suffix=suffix) # Keep simple suffix logic
            temp_path = temp_file_obj.name
            temp_file_obj.close()

            if HAS_REQUESTS:
                logger.debug("Usando 'requests' para download.")
                try:
                    response = requests.get(url, timeout=self.timeout, stream=True, verify=False)
                    response.raise_for_status()

                    content_type = response.headers.get('content-type') # <- Capture Content-Type
                    logger.debug(f"Downloaded file Content-Type (requests): {content_type}")

                    content_length = response.headers.get('content-length')
                    if content_length and int(content_length) > self.max_file_size:
                        raise Exception(f"Arquivo muito grande: {content_length} bytes (Limite: {self.max_file_size} bytes)")

                    with open(temp_path, 'wb') as f:
                         for chunk in response.iter_content(chunk_size=8192):
                             if not chunk:
                                continue
                             f.write(chunk)

                except requests.exceptions.RequestException as req_e:
                    logger.error(f"Erro requests ao baixar {url}: {req_e}")
                    logger.error(f"Traceback requests error: {traceback.format_exc()}")
                    if temp_path and os.path.exists(temp_path):
                        try:
                            os.unlink(temp_path)
                        except Exception as ul_e:
                            logger.warning(f"Failed to clean up temp file {temp_path}: {ul_e}")
                    return None # <- Return None on failure

            else:
                # Fallback para urllib
                logger.debug("Usando 'urllib' para download.")
                try:
                    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
                    req = urllib.request.Request(url, headers=headers)

                    with urllib.request.urlopen(req, timeout=self.timeout, context=ssl_context) as response:
                        content_type = response.headers.get('Content-Type') # <- Capture Content-Type (might be capitalized in urllib headers)
                        logger.debug(f"Downloaded file Content-Type (urllib): {content_type}")

                        content_length = response.headers.get('content-length')
                        if content_length and int(content_length) > self.max_file_size:
                            raise Exception(f"Arquivo muito grande: {content_length} bytes (Limite: {self.max_file_size} bytes)")

                        data = response.read()

                        with open(temp_path, 'wb') as f:
                             f.write(data)

                except urllib.error.URLError as url_e:
                     logger.error(f"Erro urllib ao baixar {url}: {url_e}")
                     logger.error(f"Traceback urllib URL error: {traceback.format_exc()}")
                     if temp_path and os.path.exists(temp_path):
                         try:
                             os.unlink(temp_path)
                         except Exception as ul_e:
                             logger.warning(f"Failed to clean up temp file {temp_path}: {ul_e}")
                     return None # <- Return None on failure
                except Exception as urll_e: # Catch other potential urllib errors
                     logger.error(f"Erro inesperado urllib ao baixar {url}: {urll_e}")
                     logger.error(f"Traceback urllib other error: {traceback.format_exc()}")
                     if temp_path and os.path.exists(temp_path):
                         try:
                             os.unlink(temp_path)
                         except Exception as ul_e:
                             logger.warning(f"Failed to clean up temp file {temp_path}: {ul_e}")
                     return None # <- Return None on failure


            elapsed_time = time.time() - start_time
            logger.info(f"Arquivo baixado com sucesso para: {temp_path} (Content-Type: {content_type}) em {elapsed_time:.2f}s")
            return (temp_path, content_type) # <- Return both path and content_type

        except Exception as e:
            logger.error(f"Erro geral fatal ao baixar arquivo {url}: {e}")
            logger.error(f"Traceback geral fatal download error: {traceback.format_exc()}")
            if temp_path and os.path.exists(temp_path):
                try: os.unlink(temp_path)
                except Exception as unlink_e: logger.warning(f"Could not delete temp file {temp_path} after download failure: {unlink_e}")
            return None # <- Always return None on failure

    async def read_file(self, file_path: str, max_pages: int = 10, summary_only: bool = False) -> Dict[str, Any]:
        """
        Lê um arquivo e retorna seu conteúdo estruturado
        """
        temp_file = None
        actual_path = file_path
        is_url = self.is_url(file_path)
        content_type_from_download = None # <- Variable to store content type if downloaded

        try:
            logger.info(f"Processando: {file_path}")
            start_time = time.time()

            # Se for URL, baixa primeiro
            if is_url:
                logger.info("URL detectada, baixando arquivo...")
                download_result = await self.download_file(file_path) # <- Call download_file
                if not download_result: # Check if download was successful
                    return {"error": f"Falha ao baixar arquivo da URL: {file_path}"}
                
                temp_file, content_type_from_download = download_result # <- Unpack the result
                actual_path = temp_file

            # Check if the file exists at the actual path
            if not os.path.exists(actual_path):
                error_msg = f"Arquivo local não encontrado após processamento de caminho/download: {actual_path}"
                logger.error(error_msg)
                # Add fallback if it was supposed to be a URL but download failed early before temp_file was set?
                # No, the download_file handles temp_path cleanup and returns None if it fails.
                return {"error": error_msg}

            file_path_obj = Path(actual_path)
            extension = file_path_obj.suffix.lower() # Get suffix from temp file path

            # --- New Logic: Use Content-Type if available and URL ---
            reader_method = None
            determined_type_source = "extension" # For logging/info

            if is_url and content_type_from_download:
                 # Use lowercase for comparison
                 lower_content_type = content_type_from_download.lower()
                 logger.debug(f"Checking Content-Type '{content_type_from_download}' ({lower_content_type}) from URL source.")

                 if 'text/html' in lower_content_type:
                      # Force HTML reader if Content-Type is html
                      if HAS_BEAUTIFULSOUP or '.html' in self.supported_extensions: # Ensure HTML reading is actually supported
                            reader_method = self._read_html
                            determined_type_source = "Content-Type (HTML)"
                            logger.info(f"Determined type as HTML from Content-Type. Using {reader_method.__name__}.")
                      else:
                            logger.warning("Content-Type is HTML but HTML parsing library not available. Falling back.")

                 elif 'application/json' in lower_content_type or 'text/json' in lower_content_type:
                      if '.json' in self.supported_extensions:
                            reader_method = self._read_json
                            determined_type_source = "Content-Type (JSON)"
                            logger.info(f"Determined type as JSON from Content-Type. Using {reader_method.__name__}.")
                      else:
                            logger.warning("Content-Type is JSON but JSON reader not available. Falling back.")
                 # Add other Content-Type mappings if needed (e.g., application/pdf, text/csv)
                 # For simplicity, we rely on extension for others for now, but Content-Type mapping is more robust for URLs.
                 # A more comprehensive approach would map common Content-Types to reader methods here.

            # If Content-Type didn't give us a reader, try extension/mimetypes fallback
            if reader_method is None:
                 # If URL and no extension, try guessing from MIME type of the downloaded file content?
                 # mimetypes.guess_type on actual_path might still not be reliable without a suffix.
                 # The initial suffix guessing in download_file was already simple.
                 # Let's stick to the extension logic, knowing Content-Type is the primary override for URLs.
                 if extension in self.supported_extensions:
                     reader_method = self.supported_extensions[extension]
                     determined_type_source = "extension"
                     logger.info(f"Determined type as {extension} from extension. Using {reader_method.__name__}.")
                 # Note: Fallback to text is handled AFTER file_info is created

            # Informações básicas do arquivo (use content_type if available)
            file_info = {
                "filename": file_path_obj.name,
                "original_path": file_path,
                "is_url": is_url,
                "extension": extension or "unknown", # Report the temp file extension or unknown
                "size": os.path.getsize(actual_path),
                # Use Content-Type from download if available, otherwise guess from path
                "mime_type": content_type_from_download or mimetypes.guess_type(actual_path)[0]
            }
            logger.info(f"Informações do arquivo: {file_info}")
            logger.info(f"Tipo de leitura determinado por: {determined_type_source}")


            # Check if a specific reader method was determined
            if reader_method:
                logger.info(f"Usando método de leitura: {reader_method.__name__}")
                # Passa parâmetros especiais para PDF
                if extension == '.pdf' or (is_url and content_type_from_download and 'application/pdf' in content_type_from_download.lower()): # Also check CT for PDF
                     content = await reader_method(actual_path, max_pages, summary_only)
                else:
                    content = await reader_method(actual_path)

            else:
                # --- Original Fallback to Text ---
                # If extension is still missing or not supported AND Content-Type didn't help, try reading as text
                error_msg = f"Tipo de arquivo não suportado via extensão ({extension}) ou Content-Type ({content_type_from_download}). Tentando ler como texto."
                logger.warning(error_msg)
                file_info["warning"] = f"Unsupported file type ({extension}/{content_type_from_download}), read as plain text." # Add warning to file_info

                # Attempt to read as generic text
                content = await self._read_text(actual_path)
                if "error" in content:
                    # If reading as text also failed
                    return {"error": f"Tipo de arquivo não suportado e falha ao ler como texto: {content['error']}", "file_info": file_info}
                # If reading as text was successful, content is ready


            elapsed_time = time.time() - start_time
            logger.info(f"Conteúdo lido em {elapsed_time:.2f}s")

            # Combine results
            result = {
                "file_info": file_info,
                "content": content, # Content from the chosen reader method (_read_html or _read_text fallback)
                "processing_time": elapsed_time,
                "success": "error" not in content # Success is true unless the content part itself reported an error
            }

            # Add specific warning if fallback occurred
            if 'warning' in file_info:
                 result["warning"] = file_info["warning"]
                 del file_info["warning"] # Clean up file_info

            return result

        except Exception as e:
            # ... (restante do código de tratamento de erro geral) ...
            error_msg = f"Erro geral ao processar arquivo {file_path}: {str(e)}"
            logger.error(error_msg)
            logger.error(f"Traceback completo: {traceback.format_exc()}")
            return {"error": error_msg, "traceback": traceback.format_exc()}
        finally:
            # Remove arquivo temporário se existir
            if temp_file and os.path.exists(temp_file):
                try:
                    os.unlink(temp_file)
                    logger.info(f"Arquivo temporário removido: {temp_file}")
                except Exception as unlink_error:
                     logger.warning(f"Falha ao remover arquivo temporário {temp_file}: {unlink_error}")


    async def _read_text(self, file_path: str) -> Dict[str, Any]:
        """Lê arquivos de texto simples"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        full_content = None

        for encoding in encodings:
            try:
                logger.debug(f"Trying encoding: {encoding}")
                with open(file_path, 'r', encoding=encoding) as f:
                    full_content = f.read()
                logger.debug(f"Successfully read with encoding: {encoding}")
                break # Stop on first successful encoding
            except UnicodeDecodeError:
                logger.debug(f"Encoding failed for {encoding}")
                continue # Try next encoding
            except Exception as e:
                logger.error(f"Unexpected error reading text file with {encoding}: {e}")
                # Depending on the error, you might want to continue or break
                continue # Continue trying other encodings

        if full_content is None:
             error_msg = "Não foi possível decodificar o arquivo com encodings suportados"
             logger.error(error_msg)
             return {"error": error_msg}

        # Limit content size for output
        content_limit = 10000
        content_truncated = len(full_content) > content_limit
        limited_content = full_content[:content_limit] if content_truncated else full_content

        return {
            "type": "text",
            "text": limited_content,
            "text_truncated": content_truncated,
            "full_length": len(full_content),
            "encoding": encoding, # This will be the last successful encoding tried
            "lines": len(full_content.splitlines()),
            "characters": len(full_content)
        }


    async def _read_html(self, file_path: str) -> Dict[str, Any]:
        """Lê arquivos HTML"""
        try:
            logger.debug(f"Reading HTML file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Limit content size for raw_html output
            raw_html_limit = 20000
            html_truncated = len(html_content) > raw_html_limit
            limited_html = html_content[:raw_html_limit] if html_truncated else html_content


            result = {
                "type": "html",
                "raw_html": limited_html,
                "html_truncated": html_truncated,
                "full_html_length": len(html_content)
            }

            if HAS_BEAUTIFULSOUP and BeautifulSoup:
                logger.debug("Using BeautifulSoup for HTML parsing.")
                try:
                    soup = BeautifulSoup(html_content, 'html.parser')
                    text_content = soup.get_text(separator='\n', strip=True) # Using separator for better text blocks

                    # Limit text_content output
                    text_limit = 5000
                    text_truncated_output = len(text_content) > text_limit
                    limited_text_content = text_content[:text_limit] if text_truncated_output else text_content

                    # Limit links and images lists
                    links_limit = 20
                    images_limit = 20

                    links = [a.get('href') for a in soup.find_all('a', href=True)]
                    images = [img.get('src') for img in soup.find_all('img', src=True)]

                    result.update({
                        "title": soup.title.string if soup.title else None,
                        "text_content": limited_text_content,
                        "text_content_truncated": text_truncated_output,
                        "full_text_length": len(text_content),
                        "links": links[:links_limit],
                        "links_truncated": len(links) > links_limit,
                        "total_links": len(links),
                        "images": images[:images_limit],
                        "images_truncated": len(images) > images_limit,
                        "total_images": len(images),
                        "parsing_method": "BeautifulSoup"
                    })
                except Exception as parse_error:
                     logger.warning(f"BeautifulSoup parsing failed: {parse_error}")
                     result["parsing_error"] = str(parse_error)
                     result["parsing_method"] = "Failed_BeautifulSoup"
                     # Keep raw_html and other info
            else:
                logger.warning("BeautifulSoup not available. Only raw HTML will be returned.")
                result["parsing_method"] = "Raw_Text"


            return result

        except Exception as e:
            error_msg = f"Erro geral ao processar HTML: {str(e)}"
            logger.error(error_msg)
            return {"error": error_msg}


    async def _read_xml(self, file_path: str) -> Dict[str, Any]:
        """Lê arquivos XML"""
        try:
            logger.debug(f"Reading XML file: {file_path}")
            if not HAS_XML or not ET:
                logger.warning("xml.etree.ElementTree not available, reading as raw text.")
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    content_limit = 10000
                    content_truncated = len(content) > content_limit
                    limited_content = content[:content_limit] if content_truncated else content
                    return {
                        "type": "xml",
                        "raw_xml": limited_content,
                        "xml_truncated": content_truncated,
                        "full_xml_length": len(content),
                        "error": "xml.etree.ElementTree not available for parsing, returned raw text."
                    }

            # Existing XML parsing logic...
            try:
                tree = ET.parse(file_path)
                root = tree.getroot()

                def xml_to_dict(element, depth=0, max_depth=3, child_limit_per_level=50):
                    if depth > max_depth:
                        return {"truncated": True, "tag": element.tag}

                    result = {}
                    if element.text and element.text.strip():
                        text = element.text.strip()
                        text_limit = 1000
                        text_truncated = len(text) > text_limit
                        result['text'] = text[:text_limit] if text_truncated else text
                        if text_truncated:
                             result['text_truncated'] = True

                    if element.attrib:
                        result['attributes'] = element.attrib

                    children = {}
                    num_children_processed = 0
                    for child in element:
                        if num_children_processed >= child_limit_per_level:
                             children['_truncated_children'] = f"... and {len(element) - num_children_processed} more children at this level"
                             break

                        child_data = xml_to_dict(child, depth + 1, max_depth, child_limit_per_level)
                        if child.tag in children:
                            if not isinstance(children[child.tag], list):
                                children[child.tag] = [children[child.tag]]
                            children[child.tag].append(child_data)
                        else:
                            children[child.tag] = child_data
                        num_children_processed += 1

                    if children:
                        result['children'] = children

                    return result

                return {
                    "type": "xml",
                    "root_tag": root.tag,
                    "structure": xml_to_dict(root),
                    "parsing_method": "ElementTree"
                }
            except Exception as parse_error:
                logger.warning(f"XML parsing failed with ElementTree: {parse_error}. Reading as raw text.")
                # Fallback to reading as text on parse error
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    content_limit = 10000
                    content_truncated = len(content) > content_limit
                    limited_content = content[:content_limit] if content_truncated else content
                    return {"type": "xml", "raw_xml": limited_content, "xml_truncated": content_truncated, "full_xml_length": len(content), "parsing_error": str(parse_error), "parsing_method": "Raw_Text_Fallback"}


        except Exception as e:
            error_msg = f"Erro geral ao processar XML: {str(e)}"
            logger.error(error_msg)
            return {"error": error_msg}


    async def _read_json(self, file_path: str) -> Dict[str, Any]:
        """Lê arquivos JSON"""
        try:
            logger.debug(f"Reading JSON file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)

            return {
                "type": "json",
                # Do not return the full 'data' object directly for potentially large JSON
                # "data": json_data, # Consider removing or truncating this
                "structure": self._analyze_json_structure(json_data),
                "size_info": {
                    "total_keys_root": len(json_data) if isinstance(json_data, dict) else None,
                    "total_items_root": len(json_data) if isinstance(json_data, list) else None,
                    "json_size_bytes": os.path.getsize(file_path)
                }
            }

        except Exception as e:
            error_msg = f"Erro ao processar JSON: {str(e)}"
            logger.error(error_msg)
            return {"error": error_msg}

    async def _read_csv(self, file_path: str) -> Dict[str, Any]:
        """Lê arquivos CSV"""
        try:
            logger.debug(f"Reading CSV file: {file_path}")
            if not HAS_PANDAS or not pd:
                logger.warning("pandas not available, using basic csv.")
                # Fallback para CSV simples sem pandas
                import csv
                try:
                    # Try multiple delimiters and sniffer
                    data = []
                    total_rows = 0
                    columns = []
                    encoding = 'utf-8' # Default
                    # Try guessing encoding first
                    try:
                        import chardet
                        with open(file_path, 'rb') as f:
                             raw_data = f.read(10000) # Read sample
                             detection = chardet.detect(raw_data)
                             encoding = detection['encoding'] or 'utf-8'
                             logger.debug(f"Guessed encoding: {encoding} with confidence {detection['confidence']:.2f}")
                    except ImportError:
                         logger.debug("chardet not available for encoding detection.")
                    except Exception as enc_e:
                         logger.warning(f"Encoding detection failed: {enc_e}. Using utf-8.")
                         encoding = 'utf-8'


                    with open(file_path, 'r', encoding=encoding) as f:
                         # Peek to guess delimiter using sniffer if available
                         try:
                             # Read a sample to sniff
                             sample = f.read(1024 * 10) # Read up to 10KB for sniffing
                             f.seek(0) # Reset file pointer
                             dialect = csv.Sniffer().sniff(sample, delimiters=',;\t|')
                             reader = csv.DictReader(f, dialect=dialect)
                             logger.debug(f"Sniffed delimiter: '{dialect.delimiter}'")
                         except Exception as sniff_error:
                            logger.warning(f"CSV Sniffer failed: {sniff_error}. Using default delimiter ','.")
                            f.seek(0) # Reset file pointer
                            reader = csv.DictReader(f)

                         # Read data and limit rows
                         sample_size = 100
                         for i, row in enumerate(reader):
                             if i < sample_size:
                                 data.append(row)
                             total_rows += 1
                         columns = list(data[0].keys()) if data else list(reader.fieldnames if reader.fieldnames else []) # Get columns from first row or fieldnames


                    # Limita número de linhas retornadas
                    sample_data = data # Already limited during iteration
                    data_truncated = total_rows > sample_size

                    return {
                        "type": "csv",
                        "data": sample_data,
                        "data_truncated": data_truncated,
                        "total_rows": total_rows,
                        "sample_rows": len(sample_data),
                        "columns": columns,
                        "reading_method": "basic_csv"
                    }

                except Exception as basic_csv_error:
                     error_msg = f"Erro ao processar CSV com csv básico: {str(basic_csv_error)}"
                     logger.error(error_msg)
                     return {"error": error_msg, "reading_method": "basic_csv_failed"}

            # Usando pandas
            logger.debug("Using pandas to read CSV.")
            try:
                df = pd.read_csv(file_path)

                # For CSVs grandes, retorna apenas amostra
                sample_size = min(100, len(df))
                sample_df = df.head(sample_size)

                return {
                    "type": "csv",
                    "data": sample_df.to_dict('records'),
                    "data_truncated": len(df) > sample_size,
                    "total_rows": len(df),
                    "sample_rows": sample_size,
                    "columns": df.columns.tolist(),
                    "dtypes": df.dtypes.astype(str).to_dict(),
                    "summary": df.describe().to_dict() if df.select_dtypes(include='number').shape[1] > 0 else None,
                    "reading_method": "pandas"
                }
            except Exception as pandas_csv_error:
                error_msg = f"Erro ao processar CSV com pandas: {str(pandas_csv_error)}. Fallback to basic csv failed previously."
                logger.error(error_msg)
                return {"error": error_msg, "reading_method": "pandas_failed"}


        except Exception as e:
            error_msg = f"Erro geral ao processar CSV: {str(e)}"
            logger.error(error_msg)
            return {"error": error_msg}


    async def _read_excel(self, file_path: str) -> Dict[str, Any]:
        """Lê arquivos Excel"""
        try:
            logger.debug(f"Reading Excel file: {file_path}")
            if not HAS_PANDAS or not pd:
                return {"error": "Pandas não está instalado para leitura de Excel"}

            # Check if openpyxl is required for .xlsx and available
            if file_path.lower().endswith('.xlsx') and not HAS_OPENPYXL:
                 return {"error": "openpyxl não está instalado, necessário para arquivos .xlsx"}

            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}

            sheets_limit = 5
            sheets_truncated_output = len(excel_file.sheet_names) > sheets_limit

            for sheet_name in excel_file.sheet_names[:sheets_limit]:  # Máximo 'sheets_limit' abas
                logger.debug(f"Reading sheet: {sheet_name}")
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    sample_size = min(50, len(df))  # Máximo 50 linhas por aba
                    sample_df = df.head(sample_size)

                    sheets_data[sheet_name] = {
                        "data": sample_df.to_dict('records'),
                        "data_truncated": len(df) > sample_size,
                        "total_rows": len(df),
                        "sample_rows": sample_size,
                        "columns": df.columns.tolist(),
                        "dtypes": df.dtypes.astype(str).to_dict()
                    }
                except Exception as sheet_error:
                     logger.warning(f"Erro ao ler aba '{sheet_name}': {sheet_error}")
                     sheets_data[sheet_name] = {"error": f"Erro ao ler aba: {str(sheet_error)}"}


            return {
                "type": "excel",
                "sheets": list(excel_file.sheet_names),
                "sheets_truncated": sheets_truncated_output,
                "total_sheets": len(excel_file.sheet_names),
                "data": sheets_data,
                 "reading_method": "pandas"
            }

        except Exception as e:
            error_msg = f"Erro geral ao processar Excel: {str(e)}"
            logger.error(error_msg)
            return {"error": error_msg}


    async def _read_pdf(self, file_path: str, max_pages: int = 10, summary_only: bool = False) -> Dict[str, Any]:
        """
        Lê arquivos PDF com controle de páginas e modo resumo

        Args:
            file_path: Caminho do arquivo PDF
            max_pages: Número máximo de páginas a processar (se summary_only é False)
            summary_only: Se True, retorna apenas informações básicas e primeiras 2 páginas
        """
        try:
            logger.info(f"Processando PDF: {file_path} (max_pages={max_pages}, summary_only={summary_only})")

            result = {"type": "pdf", "pages": [], "processing_limited": False, "extraction_method": "None"}
            total_pages = 0
            processed_pages_count = 0
            page_text_limit = 5000 # Limit character count per page text output

            # Determine how many pages to attempt to process
            pages_to_attempt = max_pages
            if summary_only:
                 pages_to_attempt = min(2, max_pages) # Limit to first 2 pages in summary mode, but respect max_pages limit too

            # Tenta usar pdfplumber primeiro
            if HAS_PDFPLUMBER and pdfplumber:
                logger.info("Attempting pdfplumber extraction.")
                try:
                    with pdfplumber.open(file_path) as pdf:
                        total_pages = len(pdf.pages)
                        logger.debug(f"PDF has {total_pages} pages (pdfplumber).")

                        # Determine final number of pages to process based on limits
                        pages_to_process = min(pages_to_attempt, total_pages)
                        result["processing_limited"] = pages_to_process < total_pages or summary_only # Limited if we process fewer than total or if summary_only was requested

                        for i in range(pages_to_process):
                            try:
                                logger.debug(f"Extracting page {i + 1}/{total_pages} with pdfplumber...")
                                page = pdf.pages[i]
                                page_text = page.extract_text() or ""

                                # Apply text content limit
                                text_truncated_output = len(page_text) > page_text_limit
                                limited_page_text = page_text[:page_text_limit] if text_truncated_output else page_text

                                result["pages"].append({
                                    "page_number": i + 1,
                                    "text": limited_page_text,
                                    "text_truncated": text_truncated_output,
                                    "full_text_length_page": len(page_text),
                                    "extraction_status": "success"
                                })
                                processed_pages_count += 1
                                logger.debug(f"Page {i + 1}: extracted {len(page_text)} chars (limited to {len(limited_page_text)})")

                            except Exception as page_error:
                                logger.warning(f"Error extracting page {i + 1} with pdfplumber: {page_error}")
                                result["pages"].append({
                                    "page_number": i + 1,
                                    "text": "",
                                    "text_truncated": False,
                                    "full_text_length_page": 0,
                                    "extraction_status": "error",
                                    "error": f"pdfplumber page error: {str(page_error)}"
                                })

                        result["total_pages"] = total_pages
                        result["processed_pages"] = processed_pages_count
                        result["extraction_method"] = "pdfplumber"
                        # Calculate full text from processed pages
                        result["full_text"] = "\n---\n".join([p.get("text", "") for p in result["pages"]])


                        logger.info(f"PDF processed with pdfplumber. Processed {processed_pages_count}/{total_pages} pages.")

                        # Add summary info if requested
                        if summary_only:
                             total_chars_processed = sum(p.get("full_text_length_page", 0) for p in result["pages"])
                             result["summary"] = {
                                "processed_character_count": total_chars_processed,
                                "estimated_total_character_count": int(total_chars_processed * (total_pages / processed_pages_count)) if processed_pages_count > 0 else 0,
                                "contains_text": total_chars_processed > 0 # Simple check if any text was extracted
                             }
                             # Remove page texts from summary output to reduce size
                             for page in result["pages"]:
                                 if "text" in page: del page["text"]
                                 if "text_truncated" in page: del page["text_truncated"]
                                 # Keep char counts
                             if "full_text" in result: del result["full_text"] # Remove combined text

                        return result # Return result even if partial success with pdfplumber


                except Exception as pdfplumber_error:
                    logger.error(f"Critical error with pdfplumber processing {file_path}: {pdfplumber_error}")
                    # traceback.print_exc() # Optional: Print traceback for critical errors
                    result["pdfplumber_error"] = str(pdfplumber_error)
                    # Continue to try PyPDF2


            # Fallback for text extraction using PyPDF2
            if HAS_PYPDF2 and PyPDF2:
                logger.info("pdfplumber failed or not available. Attempting PyPDF2 extraction.")
                # If pdfplumber was tried but failed critically or partially,
                # clear the previous result pages to retry with PyPDF2 if we want a clean PyPDF2 result.
                # Or, keep the partial result and just add PyPDF2 as a note of failure?
                # Let's clear and retry for a consistent output structure from one library.
                if result.get("pages"): # If pdfplumber added any pages (even empty ones with errors)
                     logger.warning("Clearing previous partial pdfplumber result to retry with PyPDF2.")
                     result["pages"] = []
                     result["total_pages"] = 0
                     result["processed_pages"] = 0
                     result["full_text"] = ""
                     result["extraction_method"] = "None" # Reset method


                try:
                    with open(file_path, 'rb') as f:
                        # Handle potential different PyPDF2 versions
                        try:
                             # PyPDF2 v3+
                             reader = PyPDF2.PdfReader(f)
                        except AttributeError:
                             # PyPDF2 v < 3
                             logger.warning("PyPDF2.PdfReader not found, falling back to PdfFileReader.")
                             reader = PyPDF2.PdfFileReader(f)

                        total_pages = len(reader.pages)
                        logger.debug(f"PDF has {total_pages} pages (PyPDF2).")

                        # Determine final number of pages to process based on limits
                        pages_to_process = min(pages_to_attempt, total_pages)
                        result["processing_limited"] = pages_to_process < total_pages or summary_only

                        processed_pages_count = 0
                        for i in range(pages_to_process):
                            try:
                                logger.debug(f"Extracting page {i + 1}/{total_pages} with PyPDF2...")
                                page = reader.pages[i] # Access pages by index
                                # Use extract_text() for newer PyPDF2, extractText() for older
                                try:
                                     page_text = page.extract_text() or ""
                                except AttributeError:
                                     logger.warning("PyPDF2 page.extract_text() not found, falling back to page.extractText()")
                                     page_text = page.extractText() or "" # Old method name


                                # Apply text content limit
                                text_truncated_output = len(page_text) > page_text_limit
                                limited_page_text = page_text[:page_text_limit] if text_truncated_output else page_text


                                result["pages"].append({
                                    "page_number": i + 1,
                                    "text": limited_page_text,
                                    "text_truncated": text_truncated_output,
                                    "full_text_length_page": len(page_text),
                                    "extraction_status": "success"
                                })
                                processed_pages_count += 1
                                logger.debug(f"Page {i + 1}: extracted {len(page_text)} chars (limited to {len(limited_page_text)})")


                            except Exception as page_error:
                                logger.warning(f"Error extracting page {i + 1} with PyPDF2: {page_error}")
                                result["pages"].append({
                                    "page_number": i + 1,
                                     "text": "",
                                     "text_truncated": False,
                                     "full_text_length_page": 0,
                                     "extraction_status": "error",
                                     "error": f"PyPDF2 page error: {str(page_error)}"
                                })

                        result["total_pages"] = total_pages
                        result["processed_pages"] = processed_pages_count
                        result["extraction_method"] = "PyPDF2"
                        result["full_text"] = "\n---\n".join([p.get("text", "") for p in result["pages"]])

                        logger.info(f"PDF processed with PyPDF2. Processed {processed_pages_count}/{total_pages} pages.")

                        # Add summary info if requested
                        if summary_only:
                             total_chars_processed = sum(p.get("full_text_length_page", 0) for p in result["pages"])
                             result["summary"] = {
                                "processed_character_count": total_chars_processed,
                                "estimated_total_character_count": int(total_chars_processed * (total_pages / processed_pages_count)) if processed_pages_count > 0 else 0,
                                "contains_text": total_chars_processed > 0
                             }
                             # Remove page texts from summary output
                             for page in result["pages"]:
                                 if "text" in page: del page["text"]
                                 if "text_truncated" in page: del page["text_truncated"]
                             if "full_text" in result: del result["full_text"]

                        return result # Return result even if partial success with PyPDF2


                except Exception as pypdf2_error:
                    logger.error(f"Critical error with PyPDF2 processing {file_path}: {pypdf2_error}")
                    # traceback.print_exc() # Optional: Print traceback
                    result["pypdf2_error"] = str(pypdf2_error)
                    # Fall through to final error if both failed


            # If reached here, both libraries failed critically or weren't available
            if not result.get("pages"): # If no pages were successfully processed by either library
                 error_msg = "Nenhuma biblioteca PDF disponível ou ambas falharam na extração."
                 details = f"pdfplumber available: {HAS_PDFPLUMBER}, PyPDF2 available: {HAS_PYPDF2}. "
                 details += f"Last pdfplumber critical error: {result.get('pdfplumber_error', 'None')}. "
                 details += f"Last PyPDF2 critical error: {result.get('pypdf2_error', 'None')}."
                 logger.error(error_msg + " " + details)
                 return {"error": error_msg, "details": details}
            else:
                 # If some pages were extracted but errors occurred or limits were hit, return the partial result
                 logger.warning("PDF processing finished with partial results or errors on some pages.")
                 return result


        except Exception as e:
            # Catch any unexpected errors during the overall PDF handling flow
            logger.error(f"Erro geral inesperado ao processar PDF: {e}")
            logger.error(f"Traceback completo: {traceback.format_exc()}")
            return {"error": f"Erro geral ao processar PDF: {str(e)}", "traceback": traceback.format_exc()}


    async def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """Lê arquivos DOCX"""
        try:
            logger.debug(f"Reading DOCX file: {file_path}")
            if not HAS_DOCX or not Document:
                return {"error": "python-docx não está instalado ou não disponível."}

            doc = Document(file_path)

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Limit number of paragraphs for output
            paragraph_limit = 100
            paragraphs_truncated_output = len(paragraphs) > paragraph_limit
            limited_paragraphs = paragraphs[:paragraph_limit] if paragraphs_truncated_output else paragraphs


            tables = []
            table_limit = 5
            tables_truncated_output = len(doc.tables) > table_limit

            for i, table in enumerate(doc.tables):
                if i >= table_limit:  # Maximum 'table_limit' tables
                    break
                table_data = []
                # Limit rows and cells within tables if needed? (Optional, can add later)
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Combine paragraphs for full text, apply limit
            full_text = "\n".join(paragraphs) # Use all paragraphs for full_text before truncating
            full_text_limit = 10000
            full_text_truncated_output = len(full_text) > full_text_limit
            limited_full_text = full_text[:full_text_limit] + "...[texto truncado]" if full_text_truncated_output else full_text


            return {
                "type": "docx",
                "paragraphs": limited_paragraphs,
                "paragraphs_truncated": paragraphs_truncated_output,
                "total_paragraphs": len(paragraphs),
                "tables": tables,
                "tables_truncated": tables_truncated_output,
                "total_tables": len(doc.tables),
                "full_text": limited_full_text,
                "full_text_truncated": full_text_truncated_output,
                "full_text_length": len(full_text)
            }

        except Exception as e:
            error_msg = f"Erro ao processar DOCX: {str(e)}"
            logger.error(error_msg)
            return {"error": error_msg}


    def _analyze_json_structure(self, data: Any, max_depth: int = 3) -> Dict[str, Any]:
        """Analisa a estrutura de dados JSON"""
        if max_depth <= 0:
            return {"type": type(data).__name__, "truncated": True}

        if isinstance(data, dict):
            keys = list(data.keys())
            keys_limit = 20
            keys_truncated_output = len(keys) > keys_limit
            limited_keys = keys[:keys_limit] if keys_truncated_output else keys

            items_limit_for_structure_analysis = 5 # How many key-value pairs to show structure for at current depth

            return {
                "type": "object",
                "keys": limited_keys,
                "keys_truncated": keys_truncated_output,
                "total_keys": len(keys),
                "structure": {k: self._analyze_json_structure(v, max_depth - 1)
                            for k, v in list(data.items())[:items_limit_for_structure_analysis]},
                "structure_items_limited": len(data) > items_limit_for_structure_analysis
            }
        elif isinstance(data, list):
            list_length = len(data)
            return {
                "type": "array",
                "length": list_length,
                "sample_structure": self._analyze_json_structure(data[0], max_depth - 1) if data else None,
                "all_items_same_structure": True # Assumption, could add checks for homogeneity
            }
        else:
            return {"type": type(data).__name__}


# Instância global do leitor de arquivos
# Only create if MCP is potentially going to run
if HAS_MCP:
    file_reader_instance = FileReader()
    logger.info("FileReader instance created.")
else:
    file_reader_instance = None
    logger.warning("FileReader instance not created because MCP is not available.")


# Register tools only if MCP instance was created
if mcp and file_reader_instance:

    @mcp.tool()
    async def read_file(
        file_path: Optional[str] = None,
        file_content: Optional[str] = None,
        filename: Optional[str] = None,
        max_pages: int = 10,
        summary_only: bool = False
    ) -> Dict[str, Any]:
        """
        Lê e extrai conteúdo de arquivos de diversos formatos, incluindo URLs
        
        Args:
            file_path: Caminho local ou URL do arquivo
            file_content: Conteúdo do arquivo em base64 (alternativa ao file_path)
            filename: Nome do arquivo (necessário se usar file_content)
            max_pages: Número máximo de páginas para PDFs (default: 10). Ignorado para outros tipos.
            summary_only: Se True, retorna apenas resumo para arquivos grandes (default: False). Aplica-se a PDFs (primeiras 2 páginas) e outros tipos (amostras).
        """

        logger.info(f"Tool 'read_file' called with path={file_path}, filename={filename}, content_provided={file_content is not None}, max_pages={max_pages}, summary_only={summary_only}")

        if file_content and filename:
            logger.debug("Processing file from base64 content.")
            temp_path = None
            try:
                content_bytes = base64.b64decode(file_content)

                # Create temp file with suffix from filename
                suffix = Path(filename).suffix
                # Add a basic guess if no suffix
                if not suffix:
                    mime = mimetypes.guess_type(filename)[0]
                    if mime:
                        guessed_suffix = mimetypes.guess_extension(mime)
                        if guessed_suffix:
                            suffix = guessed_suffix
                            logger.debug(f"Guessed suffix '{suffix}' from mime '{mime}' for filename '{filename}'")
                        else:
                             suffix = ""
                    else:
                        suffix = "" # No suffix guess

                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                    tmp_file.write(content_bytes)
                    temp_path = tmp_file.name

                logger.info(f"Created temp file from base64: {temp_path}")

                try:
                    # Call the core file reading logic
                    result = await file_reader_instance.read_file(temp_path, max_pages, summary_only)
                finally:
                    # Ensure temp file is removed
                    if temp_path and os.path.exists(temp_path):
                        try:
                            os.unlink(temp_path)
                            logger.debug(f"Removed temp file: {temp_path}")
                        except Exception as unlink_error:
                            logger.warning(f"Failed to remove temp file {temp_path}: {unlink_error}")


            except Exception as e:
                logger.error(f"Error processing base64 file: {e}")
                result = {"error": f"Erro ao processar arquivo base64: {str(e)}", "traceback": traceback.format_exc()}

        elif file_path:
            logger.debug(f"Processing file from path: {file_path}")
            # Call the core file reading logic
            result = await file_reader_instance.read_file(file_path, max_pages, summary_only)
        else:
            error_msg = "É necessário fornecer file_path (local ou URL) ou file_content + filename"
            logger.error(error_msg)
            result = {"error": error_msg}

        logger.info(f"Tool 'read_file' finished. Success: {result.get('success', False)}, Error: {'error' in result}")
        return result

    @mcp.tool()
    async def read_url(url: str, max_pages: int = 5, summary_only: bool = True) -> Dict[str, Any]:
        """
        Versão otimizada para ler arquivos de URLs.
        Baixa o arquivo temporariamente e extrai seu conteúdo.
        Usa defaults mais restritivos para URLs (max_pages=5, summary_only=True).

        Args:
            url: A URL do arquivo a ser lido.
            max_pages: Número máximo de páginas para PDFs (default: 5). Ignorado para outros tipos.
            summary_only: Se True, retorna apenas resumo para arquivos grandes (default: True). Aplica-se a PDFs (primeiras 2 páginas) e outros tipos (amostras).
        """
        logger.info(f"Tool 'read_url' called with url={url}, max_pages={max_pages}, summary_only={summary_only}")

        # Validate if it looks like a URL
        if not file_reader_instance.is_url(url):
             error_msg = f"Entrada não parece uma URL válida: {url}"
             logger.error(error_msg)
             return {"error": error_msg}

        # Call the main read_file logic with the URL as the path
        # The read_file method already handles URL downloading internally
        result = await file_reader_instance.read_file(url, max_pages=max_pages, summary_only=summary_only)

        logger.info(f"Tool 'read_url' finished. Success: {result.get('success', False)}, Error: {'error' in result}")
        return result


    @mcp.tool()
    async def list_supported_formats() -> Dict[str, Any]:
        """
        Lista todos os formatos de arquivo suportados e status das bibliotecas.
        """
        logger.info("Tool 'list_supported_formats' called.")
        if not file_reader_instance:
             return {"error": "FileReader instance not initialized.", "library_status": {}}

        formats = {
            "supported_extensions": list(file_reader_instance.supported_extensions.keys()), # Get live supported extensions
            "description": "Formatos suportados para leitura e extração de conteúdo (based on available libraries)",
            "library_status": {
                "pdf": {
                    "pdfplumber": HAS_PDFPLUMBER,
                    "PyPDF2": HAS_PYPDF2,
                    "notes": "Requires at least one of these for PDF."
                },
                "docx": HAS_DOCX,
                "excel": HAS_PANDAS, # pandas handles both .xlsx and .xls
                "csv": HAS_PANDAS, # pandas preferred for CSV, basic csv fallback available
                "html_parsing": HAS_BEAUTIFULSOUP, # BeautifulSoup for parsing, raw text fallback available
                "xml_parsing": HAS_XML, # ElementTree for parsing, raw text fallback available
                "url_download": HAS_REQUESTS # requests preferred for downloading, urllib fallback available
            }
        }
        logger.debug(f"Supported formats result: {formats}")
        return formats

    @mcp.tool()
    async def diagnose_pdf_libraries() -> Dict[str, Any]:
        """
        Diagnóstica o status das bibliotecas PDF e tenta importá-las para verificar.
        """
        logger.info("Tool 'diagnose_pdf_libraries' called.")
        diagnosis = {
            "pdfplumber": {
                "initial_check_available": HAS_PDFPLUMBER,
                "module_object_str": str(pdfplumber) if pdfplumber else None,
                "runtime_import_test_status": "NOT_TESTED",
                "runtime_import_error": None
            },
            "PyPDF2": {
                "initial_check_available": HAS_PYPDF2,
                "module_object_str": str(PyPDF2) if PyPDF2 else None,
                "runtime_import_test_status": "NOT_TESTED",
                 "runtime_import_error": None
            }
        }

        # Try importing again at runtime to catch specific environment errors
        try:
            import pdfplumber as test_pdfplumber
            diagnosis["pdfplumber"]["runtime_import_test_status"] = "SUCCESS"
            # Clean up the temporary import
            del test_pdfplumber
        except Exception as e:
            diagnosis["pdfplumber"]["runtime_import_test_status"] = "FAILURE"
            diagnosis["pdfplumber"]["runtime_import_error"] = str(e)
            logger.error(f"Runtime import test failed for pdfplumber: {e}")

        try:
            import PyPDF2 as test_pypdf2
            diagnosis["PyPDF2"]["runtime_import_test_status"] = "SUCCESS"
             # Clean up the temporary import
            del test_pypdf2
        except Exception as e:
            diagnosis["PyPDF2"]["runtime_import_test_status"] = "FAILURE"
            diagnosis["PyPDF2"]["runtime_import_error"] = str(e)
            logger.error(f"Runtime import test failed for PyPDF2: {e}")

        logger.debug(f"PDF diagnosis result: {diagnosis}")
        return diagnosis

else:
    logger.critical("MCP instance not created due to import errors or missing libraries. Tools will not be registered.")


if __name__ == "__main__":
    logger.info("Starting MCP File Reader Server")

    # Log final library status after init and potential removals
    if 'file_reader_instance' in locals() and file_reader_instance:
         logger.info(f"Final Supported Extensions: {list(file_reader_instance.supported_extensions.keys())}")
    else:
         logger.warning("FileReader instance not available at startup.")

    if mcp:
        # --- Modificação para SSE: Mudar o transporte para "sse" ---
        transport = "sse" 
        if transport == "stdio":
            print("Running server with stdio transport")
            mcp.run(transport="stdio")
        elif transport == "sse":
            print(f"Running server with SSE transport")
            mcp.run(transport="sse")
        else:
            raise ValueError(f"Unknown transport: {transport}")
    else:
        logger.critical("MCP server cannot run because the FastMCP instance was not created.")
        sys.exit(1)